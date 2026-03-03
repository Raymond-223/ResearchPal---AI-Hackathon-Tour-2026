"""
Export utilities for ResearchPal
Generate formatted exports in multiple formats
"""
from typing import Optional, List, Dict
from datetime import datetime
import json

class ExportManager:
    """Handle export to various formats"""

    @staticmethod
    def export_to_markdown(
        title: str,
        content: str,
        metadata: Optional[Dict] = None
    ) -> str:
        """Export content as formatted Markdown"""
        lines = []

        # Header
        lines.append(f"# {title}\n")

        # Metadata
        if metadata:
            lines.append("---")
            lines.append(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            if metadata.get("source"):
                lines.append(f"来源: {metadata['source']}")
            lines.append("---\n")

        # Content
        lines.append(content)

        return "\n".join(lines)

    @staticmethod
    def export_to_docx(
        title: str,
        content: str,
        metadata: Optional[Dict] = None
    ) -> bytes:
        """Export content as Word document"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Title
            title_para = doc.add_heading(title, level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Metadata
            if metadata:
                meta_para = doc.add_paragraph()
                meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                meta_run = meta_para.add_run(
                    f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
                )
                meta_run.font.size = Pt(10)
                meta_run.font.color.rgb = RGBColor(128, 128, 128)

                doc.add_paragraph()  # Spacing

            # Content
            for paragraph in content.split('\n\n'):
                if paragraph.strip():
                    p = doc.add_paragraph(paragraph.strip())

            # Export to bytes
            from io import BytesIO
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.read()

        except ImportError:
            raise ImportError(
                "python-docx not installed. Install with: pip install python-docx"
            )

    @staticmethod
    def format_citation(
        title: str,
        authors: List[str],
        year: int,
        journal: Optional[str] = None,
        doi: Optional[str] = None,
        style: str = "APA"
    ) -> str:
        """Format citation in various styles"""

        if style == "APA":
            # APA format: Author(s). (Year). Title. Journal, DOI.
            authors_str = ", ".join(authors[:3])
            if len(authors) > 3:
                authors_str += ", et al."

            citation = f"{authors_str} ({year}). {title}"
            if journal:
                citation += f". {journal}"
            if doi:
                citation += f". {doi}"
            return citation

        elif style == "MLA":
            # MLA format: Author(s). "Title." Journal, Year, DOI.
            authors_str = ", ".join(authors[:2])
            if len(authors) > 2:
                authors_str += ", et al."

            citation = f'{authors_str}. "{title}."'
            if journal:
                citation += f" {journal},"
            citation += f" {year}"
            if doi:
                citation += f", {doi}"
            citation += "."
            return citation

        elif style == "IEEE":
            # IEEE format: [1] Author(s), "Title," Journal, Year, DOI.
            authors_str = ", ".join(authors[:3])
            if len(authors) > 3:
                authors_str += ", et al."

            citation = f'{authors_str}, "{title},"'
            if journal:
                citation += f" {journal},"
            citation += f" {year}"
            if doi:
                citation += f", {doi}"
            citation += "."
            return citation

        elif style == "Chicago":
            # Chicago format: Author(s). "Title." Journal (Year): DOI.
            authors_str = ", ".join(authors[:3])
            if len(authors) > 3:
                authors_str += ", et al."

            citation = f'{authors_str}. "{title}."'
            if journal:
                citation += f" {journal}"
            citation += f" ({year})"
            if doi:
                citation += f": {doi}"
            citation += "."
            return citation

        elif style == "GB/T 7714":
            # Chinese national standard
            authors_str = ", ".join(authors[:3])
            if len(authors) > 3:
                authors_str += ", 等."

            citation = f"{authors_str} {title}[J]"
            if journal:
                citation += f". {journal}"
            citation += f", {year}"
            if doi:
                citation += f". DOI: {doi}"
            return citation

        else:
            # Default to APA
            return ExportManager.format_citation(
                title, authors, year, journal, doi, "APA"
            )

    @staticmethod
    def export_bibliography(
        references: List[Dict],
        style: str = "APA"
    ) -> str:
        """Export formatted bibliography"""
        lines = ["# 参考文献\n"]

        for ref in references:
            citation = ExportManager.format_citation(
                title=ref.get("title", ""),
                authors=ref.get("authors", ["Unknown"]),
                year=ref.get("year", datetime.now().year),
                journal=ref.get("journal"),
                doi=ref.get("doi"),
                style=style
            )
            lines.append(f"{citation}\n")

        return "\n".join(lines)

    @staticmethod
    def save_export(
        content: str or bytes,
        filename: str,
        format: str
    ) -> str:
        """Save export to file"""
        import os

        export_dir = "exports"
        os.makedirs(export_dir, exist_ok=True)

        filepath = os.path.join(export_dir, filename)

        if isinstance(content, bytes):
            mode = "wb"
        else:
            mode = "w"
            if format in ["md", "markdown", "bib"]:
                encoding = "utf-8"

        with open(filepath, mode, encoding=encoding if 'encoding' in locals() else None) as f:
            f.write(content)

        return filepath


class ClipboardManager:
    """Manage clipboard operations"""

    @staticmethod
    def copy_to_clipboard(text: str) -> bool:
        """Copy text to system clipboard"""
        try:
            import pyperclip
            pyperclip.copy(text)
            return True
        except ImportError:
            # Fallback: generate copy command
            return False

    @staticmethod
    def generate_copy_script(text: str) -> str:
        """Generate JavaScript copy script for Gradio"""
        return f"""
        <script>
        navigator.clipboard.writeText(`{text}`).then(function() {{
            console.log('Copied to clipboard');
        }}).catch(function(err) {{
            console.error('Copy failed:', err);
        }});
        </script>
        """
